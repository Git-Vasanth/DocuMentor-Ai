from docx import Document

path = r"C:\Users\Vasanth\Desktop\Desktop\TalkDoc.docx"

doc = Document(path)

doc = Document(path)

# Iterate over all paragraphs and print their text
for i, paragraph in enumerate(doc.paragraphs):
    print(f"Paragraph {i+1}: {paragraph.text}")